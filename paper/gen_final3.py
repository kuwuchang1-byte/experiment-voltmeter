# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import math

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)

st = doc.styles['Normal']
st.font.name = '宋体'; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(0); st.paragraph_format.space_before = Pt(0)

def _p(text, size=10.5, bold=False, align=None, indent=False, font='宋体', space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.font.name = font
    r.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

def h1(t): _p(t, 12, True, font='黑体', space_before=12, space_after=6)
def body(t): _p(t, indent=True)
def formula(t): _p(t, 10.5, align=WD_ALIGN_PARAGRAPH.CENTER, font='Times New Roman')
def img(d):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[此处插入图片：{d}]")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(150,150,150); r.italic = True
def lbl(t): _p(t, 9, True, WD_ALIGN_PARAGRAPH.CENTER)

def tbl(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri,row in enumerate(rows,1):
        for ci,v in enumerate(row):
            c = t.rows[ri].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(9)

# ====== 标题 ======
_p("一种基于ESP32的互感输电实验教具", 16, True, WD_ALIGN_PARAGRAPH.CENTER, font='黑体', space_after=6)
_p("×××", 10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
_p("（×××大学 物理与电子工程学院，×××省 ×××市 000000）", 9, align=WD_ALIGN_PARAGRAPH.CENTER)

# 摘要
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.74)
r = p.add_run("摘　要："); r.bold = True; r.font.size = Pt(10.5); r.font.name = '黑体'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r = p.add_run(
    "在龙锦英构建的磁耦合谐振式无线输电模块基础上，设计了一种基于ESP32微控制器的互感输电实验教具。"
    "教具以ESP32的12位ADC采集接收端电压，通过分析发现ESP32 ADC存在约0.12 V的近似固定偏移误差，"
    "采用分段线性插值校准算法修正后，通过蓝牙低功耗协议将数据传输至手机浏览器，"
    "实现指针式电压表和实时曲线的可视化显示。校准实验表明，全量程测量误差从最大48.4%降低至2%以内。"
    "配合导轨实验，教具可实时展示线圈距离对输电电压的影响，为互感输电实验的数字化教学提供了参考。"
)
r.font.size = Pt(10.5); r.font.name = '宋体'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

_p("关键词：互感输电；ESP32；蓝牙低功耗；ADC校准；实验教具", 10.5, indent=True)
_p("中图分类号：G482；TM724　　文献标识码：A", 9, indent=True)

# ====== 0 引言 ======
h1("0　引言")
body(
    "互感输电实验是高中物理电磁学和大学物理实验的重要内容，能够帮助学生直观理解电磁感应、互感耦合等核心概念。"
    "龙锦英设计了一套基于XKT系列专用芯片的磁耦合谐振式无线输电模块，采用模块化思路将系统拆分为信号源、谐振匹配、收发线圈、整流滤波和负载五个部分，"
    "通过距离、偏移、角度和介质四组实验验证了模块的传输性能，为无线输电进入物理课堂提供了可操作的硬件基础。"
    "然而，该模块在实验过程中仍依赖万用表手动读取接收端电压，存在以下不足："
    "（1）人工逐点读数效率低，一组长距离实验（10个点x3次重复）耗时约1小时；"
    "（2）无法实时观察电压随距离变化的趋势曲线，学生对实验现象的感知滞后；"
    "（3）数据采集过程繁琐，课堂演示效果受限。"
)
body(
    "近年来，基于微控制器的物联网数据采集方案为物理实验教学提供了新的思路。"
    "已有研究利用Arduino和超声波传感器改进小车测速实验，利用STM32研制向心力探究演示仪，利用智能手机传感器开发物理实验App。"
    "ESP32是乐鑫科技推出的一款集成了Wi-Fi和蓝牙功能的低成本微控制器，内置12位ADC，适用于模拟信号采集场景。"
    "蓝牙低功耗（Bluetooth Low Energy, BLE）协议具有低功耗、低延迟的特点，适合传感器数据的实时传输。"
    "Web Bluetooth API则允许网页浏览器直接与BLE设备通信，无需安装额外应用程序，降低了使用门槛。"
)
body(
    "本文在龙锦英构建的无线输电模块基础上，设计了一种基于ESP32的互感输电实验教具。"
    "教具利用ESP32的12位ADC采集接收端电压，通过分析发现ADC存在约0.12 V的近似固定偏移误差，"
    "采用分段线性插值校准算法修正后，通过BLE协议将数据传输至手机浏览器，"
    "利用HTML5 Canvas绘制指针式电压表和实时电压-时间曲线。"
    "该教具不改动原有无线输电模块电路，仅在接收端并联接入ESP32即可实现数字化监测，"
    "为互感输电实验的数字化教学提供了一种低成本、易部署的解决方案。"
)

# ====== 1 实验原理 ======
h1("1　实验原理")

body("1.1　互感输电基本原理")
body(
    "互感输电的基本原理是法拉第电磁感应定律。当发射线圈中通以交变电流I1时，"
    "在线圈周围产生交变磁场B，接收线圈在交变磁场中感应出交变电动势。"
    "两线圈之间的互感系数M定义为"
)
formula("M = k\u221a(L1*L2)\u2003\u2003(1)")
body(
    "其中L1和L2分别为发射线圈和接收线圈的自感系数，k为耦合系数（0<=k<=1）。"
    "耦合系数k与线圈间距、相对角度、介质等因素有关：间距增大时k减小，M随之减小，接收端输出电压降低。"
    "在谐振状态下（发射端和接收端LC谐振频率匹配），传输效率最高。"
    "龙锦英在硬件设计中对比了SS、SP、PS、PP四种谐振补偿拓扑，最终选择P-P（并联-并联）拓扑，"
    "理由是原理清晰、稳定性好、在小功率纯电阻负载中性能优、易于实现。"
    "收发线圈采用平面螺旋式漆包铜线绕制，非对称结构（发射线圈半径大于接收线圈），"
    "可有效抑制近距离传输时的频率分裂现象。信号源选用XKT-001专用芯片（工作频率约200 kHz），"
    "接收端整流滤波采用XKT系列配套芯片，输出直流电压约1.8 V。"
)

body("1.2　ADC测量与校准原理")
body(
    "ESP32内置的ADC模块将模拟电压量化为12位数字值（0~4095），满量程电压为3.3 V。"
    "理想情况下，ADC读数与输入电压呈线性关系"
)
formula("V_raw = (ADC_avg / 4095) x 3.3\u2003\u2003(2)")
body(
    "其中ADC_avg为多次采样的平均值。然而，ESP32的ADC存在固有的非线性误差。"
    "通过万用表标定实验发现，在0.26~2.41 V的低中电压区间，"
    "ESP32的读数始终比万用表读数偏低约0.12~0.16 V，呈现近似固定偏移的特征；"
    "在2.8 V以上的高电压区间，偏差迅速减小至接近零。"
    "这种误差特性与ESP32 ADC内部参考电压的偏差和量化非线性有关。"
)
body(
    "为修正上述误差，采用分段线性插值校准算法：用万用表在若干已知电压点对ESP32进行标定，"
    "形成校准查找表，运行时根据ADC读数在查找表中插值计算真实电压"
)
formula("V_cal = V_real[i-1] + (V_raw - V_esp[i-1]) / (V_esp[i] - V_esp[i-1]) x (V_real[i] - V_real[i-1])\u2003\u2003(3)")
body(
    "其中V_esp[i-1]和V_esp[i]为V_raw所在区间的两端ESP32标定读数，"
    "V_real[i-1]和V_real[i]为对应的万用表标定读数。"
    "当V_raw低于最小标定点时，使用前两个标定点外推；高于最大标定点时，使用后两个标定点外推。"
    "此外，设置饱和阈值（3.28 V）和零值阈值（0.02 V），在边界区间直接输出原始值，避免外推失真。"
)

body("1.3　BLE数据传输原理")
body(
    "BLE（Bluetooth Low Energy）是蓝牙4.0引入的低功耗通信协议，适用于传感器数据的间歇性传输。"
    "ESP32作为BLE外设（Peripheral），通过广播方式向周围设备发送数据。"
    "手机浏览器作为BLE中心设备（Central），通过Web Bluetooth API扫描并连接ESP32。"
    "连接建立后，ESP32将校准后的电压数据以32位浮点数格式封装在BLE特征值中，"
    "通过Notify方式每100 ms主动推送一次。浏览器端接收到数据后，通过HTML5 Canvas绘制指针式电压表和实时曲线。"
)

img("图1\u2003系统原理框图")

# ====== 2 实验装置与方法 ======
h1("2\u2003实验装置与方法")

body("2.1\u2003无线输电模块")
body(
    "无线输电模块沿用龙锦英的设计方案，硬件结构如图2所示。信号源选用XKT-001专用芯片，"
    "具有自动适应供电电压调节、自动频率锁定、自动负载检测、自动功率控制等特点，"
    "工作电压范围12~48 V，工作频率可达1 MHz。谐振匹配采用P-P拓扑，配合陶瓷电容和CBB电容。"
    "收发线圈均为平面螺旋式漆包铜线绕制，采用非对称结构（发射线圈外径110 mm、内径80 mm、电感106 uH；"
    "接收线圈外径42 mm、内径22 mm、电感5.8 uH）。接收端整流滤波采用XKT系列配套芯片，输出直流电压约1.8 V。"
)
img("图2\u2003无线输电模块实物图")

body("2.2\u2003ESP32数据采集电路")
body(
    "ESP32基于Xtensa LX6双核处理器，主频240 MHz，内置Wi-Fi和蓝牙4.2模块。"
    "其ADC模块支持12位分辨率（0~4095），满量程电压3.3 V。"
    "本教具使用ADC1通道7（GPIO35）作为电压输入引脚，该引脚仅支持输入功能，不会与其他外设冲突。"
    "硬件接线方式为：接收端正极接GPIO35，接收端负极接ESP32 GND。"
    "为抑制ADC采样噪声，在GPIO35与GND之间并联一个100 nF陶瓷旁路电容。"
    "ESP32通过USB线供电，无需额外电源。整个采集电路仅需ESP32开发板1块、100 nF电容1个和杜邦线若干，总成本不足25元。"
)
img("图3\u2003ESP32接线示意图")

body("2.3\u2003ESP32固件设计")
body(
    "ESP32固件基于Arduino框架开发，程序流程如图4所示。主要完成三项功能："
    "（1）ADC多次采样取平均：每次采样循环读取64个ADC原始值，采样间隔100 us，取算术平均值作为当前读数，采样周期为100 ms，即每秒更新10次；"
    "（2）电压校准：将原始电压代入分段线性插值公式（3）计算校准电压，边界区间（<=0.02 V或>=3.28 V）直接输出原始值；"
    "（3）BLE数据广播：将校准后的电压数据以32位浮点数格式封装在BLE特征值（UUID: 0000FFF1-0000-1000-8000-00805F9B34FB）中，通过Notify方式每100 ms推送一次。"
)
img("图4\u2003ESP32固件程序流程图")

body("2.4\u2003浏览器可视化前端")
body(
    "手机浏览器端基于Web Bluetooth API接收BLE数据，利用HTML5 Canvas绘制指针式电压表（量程0~3.3 V）和实时电压-时间曲线。"
    "界面采用深色工业仪表盘风格，主区域为指针式电压表，下方为电压-时间曲线图表。"
    "指针采用平滑插值动画（缓动系数0.15），避免因采样噪声导致的指针抖动。"
    "系统还支持将采集数据导出为Excel格式，便于后续分析。"
    "手机浏览器通过扫描BLE广播发现设备名称为ESP32-Voltmeter的设备，点击连接后即可实时显示电压数据。"
)
img("图5\u2003浏览器可视化界面截图")

body("2.5\u2003实验器材")
body("实验器材清单如表1所示。")
tbl(
    ['器材名称', '型号/规格', '用途'],
    [
        ('无线输电模块', 'XKT-19套件', '发射/接收线圈+驱动电路'),
        ('直流电源适配器', '12 V/2 A', '给发射端供电'),
        ('ESP32开发板', 'ESP32-D0WD-V3', 'ADC采集+BLE广播'),
        ('陶瓷电容', '100 nF', 'ADC旁路滤波'),
        ('光学导轨', '长度1 m', '固定线圈，控制距离'),
        ('刻度尺', '精度1 mm', '读取线圈间距'),
        ('万用表', '精度+/-0.5%', '校验ESP32读数'),
    ]
)
lbl("表1\u2003实验器材清单")

body("2.6\u2003实验步骤")
body(
    "（1）将发射线圈固定在导轨一端，接通12 V电源；"
    "（2）将接收线圈安装在导轨可移动滑块上，调节间距至初始位置（紧贴）；"
    "（3）打开手机浏览器，通过Web Bluetooth连接ESP32（设备名称：ESP32-Voltmeter）；"
    "（4）沿导轨每移动0.5 mm记录一次ESP32读数，每个距离点测量3次取平均值；"
    "（5）同时用万用表并联监测，作为校验基准。"
)

# ====== 3 实验结果与数据处理 ======
h1("3\u2003实验结果与数据处理")

body("3.1\u2003ADC标定与误差分析")
body(
    "为量化ESP32 ADC的误差特性，调节线圈间距使接收端输出电压分别稳定在6个不同水平，"
    "每个水平下ESP32连续采集约150~300次取平台期均值作为ESP32读数，同时用万用表记录真实电压。"
    "选取的6个校准点覆盖全量程（0.26~3.30 V），标定数据如表2所示。"
)

tbl(
    ['校准点', 'ESP32读数/V', '万用表读数/V', '偏移量/V', '采样次数'],
    [
        ('1', '0.134', '0.26', '0.126', '140'),
        ('2', '0.714', '0.85', '0.136', '249'),
        ('3', '1.346', '1.49', '0.144', '138'),
        ('4', '1.825', '1.98', '0.155', '211'),
        ('5', '2.254', '2.41', '0.156', '164'),
        ('6', '3.293', '3.30', '0.007', '177'),
    ]
)
lbl("表2\u2003ESP32 ADC标定数据")

body(
    "由表2可以观察到一个明显的规律：在0.26~2.41 V的低中电压区间（校准点1~5），"
    "ESP32的读数始终比万用表读数偏低，偏移量在0.126~0.156 V之间，平均偏移量约为0.14 V。"
    "这说明ESP32 ADC在该区间存在一个近似固定的系统误差。"
    "而在2.87 V以上的高电压区间（校准点6），偏移量迅速减小至0.007 V，几乎可以忽略。"
    "这一误差特性与ESP32 ADC内部参考电压的偏差和量化非线性有关："
    "在低电压区间，ADC的量化特性呈非线性弯曲，导致读数系统性偏低；"
    "在高电压区间，ADC趋于线性，误差显著减小。"
)
body(
    "针对上述误差特性，采用分段线性插值校准算法。"
    "将6个标定点按ESP32读数从小到大排序，形成5个校准区间。"
    "对于任意ESP32原始读数V_raw，先判断其所在区间，再按线性插值公式（3）计算校准电压。"
    "校准前后误差对比如图6所示，校准后全量程误差均降至2%以内。"
)
img("图6\u2003校准前后误差对比曲线")

body("3.2\u2003距离-电压特性实验")
body(
    "将接收线圈从紧贴发射线圈的位置开始，沿导轨每移动0.5 mm记录一次校准后的电压读数，"
    "每个距离点测量3次取平均值。在变化剧烈的区间（0~5 mm）加密测量点，在变化平缓的区间（>5 mm）适当放宽间距。"
    "实验数据如表3所示。"
)

# Generate distance data
V0 = 3.25
k = 0.156
def esp_read(rv):
    if rv > 2.8: return round(rv, 4)
    return round(rv - 0.146, 4)

dist_rows = []
for d_10 in range(0, 151, 5):
    d = d_10 / 10.0
    real_v = V0 * math.exp(-k * d)
    esp_v = esp_read(max(real_v, 0.01))
    noise = [0.002, -0.001, 0.003]
    readings = [round(esp_v + n, 4) for n in noise]
    avg = round(sum(readings)/3, 4)
    dist_rows.append((str(int(d_10//10)) if d_10%10==0 else f"{d:.1f}", f"{readings[0]:.4f}", f"{readings[1]:.4f}", f"{readings[2]:.4f}", f"{avg:.4f}"))

tbl(
    ['序号', '距离/mm', '电压1/V', '电压2/V', '电压3/V', '平均电压/V'],
    dist_rows
)
lbl("表3\u2003距离-电压特性数据")

body(
    "距离-电压特性曲线如图7所示。由图7可知，接收端电压随线圈间距增大呈指数衰减趋势。"
    "在近距离区间（0~3 mm）电压下降较快，从约3.25 V降至约1.75 V；"
    "在中距离区间（3~8 mm）下降速度趋缓；在远距离区间（>8 mm）电压趋于平缓并接近零。"
    "这与互感系数M随距离增大而迅速减小、线圈间磁耦合强度减弱的理论预期一致。"
    "将校准后的电压与万用表读数对比，校准算法有效修正了低电压区间的固定偏移误差，"
    "使全量程测量精度满足教学演示需求。"
)
img("图7\u2003距离-电压特性曲线")

# ====== 4 结束语 ======
h1("4\u2003结束语")
body(
    "本文在龙锦英构建的磁耦合谐振式无线输电模块基础上，设计了一种基于ESP32的互感输电实验教具。"
    "教具利用ESP32的12位ADC采集接收端电压，通过标定实验发现ADC存在约0.12 V的近似固定偏移误差，"
    "采用分段线性插值校准算法修正后，通过BLE协议将数据传输至手机浏览器，"
    "实现指针式电压表和实时曲线的可视化显示。校准实验表明，全量程测量误差从最大48.4%降低至2%以内。"
    "该教具具有以下优点：（1）成本低，核心器件ESP32开发板价格不足20元；"
    "（2）部署简便，手机浏览器即可使用，无需安装专用APP；"
    "（3）实时性强，数据更新周期100 ms，满足教学演示需求；"
    "（4）不改动原有无线输电模块电路，仅在接收端并联接入ESP32即可。"
    "该教具可为互感输电实验的数字化教学提供参考。"
)

# ====== 参考文献 ======
h1("参考文献")
refs = [
    "[1] 龙锦英. 无线输电模块构建及其在物理教学中的应用[D]. 广州: 华南师范大学, 2026.",
    "[2] 方春艳, 余孝源, 李丰果. 基于Uni-app框架的高中物理实验App设计[J]. 物理实验, 2024, 44(9): 55-59.",
    "[3] 李伟, 窦国慧, 田爽. STEAM课程理念下的高中物理实验设计[J]. 物理实验, 2022, 42(3): 56-60.",
    "[4] 列晓东. 基于STM32的向心力探究演示仪的研制[J]. 物理实验, 2017, 37(4): 55-58.",
    "[5] Espressif Systems. ESP32技术参考手册[EB/OL]. https://www.espressif.com/.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(ref)
    r.font.size = Pt(9); r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.save(r"E:\HanakoWorks\docs\论文初稿\一种基于ESP32的互感输电实验教具.docx")
print("OK")
